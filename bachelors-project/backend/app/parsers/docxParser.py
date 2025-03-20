import docx
import re
from typing import Dict, List, Any, Tuple
import logging

logger = logging.getLogger(__name__)

def extract_sections_from_docx(docx_path: str) -> Dict[str, List[Dict[str, str]]]:
    """
    Extracts sections, subsections, and their content from a DOCX file.
    Captures all content including inputs, form fields, and minimal text.
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        Dict: A dictionary with "content" key containing a list of sections with their titles and text
    """
    try:
        doc = docx.Document(docx_path)
        extracted_data = {"content": []}
        
        # First pass: collect all paragraphs and their positions
        paragraphs_info = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            paragraphs_info.append((i, text, para))
        
        # Second pass: collect all tables and their positions
        tables_info = []
        for i, table in enumerate(doc.tables):
            table_content = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_content.append(" | ".join(row_data))
            
            table_text = "\n".join(table_content)
            # Estimate table position based on its index
            tables_info.append((i, table_text))
        
        # Third pass: identify sections and associate content
        current_section = None
        current_content = []
        sections = []
        
        for i, text, para in paragraphs_info:
            # Check if this is a main section or subsection
            if re.match(r'^\d+(\.\d+)?\s+.*', text):
                # Save previous section if exists
                if current_section:
                    sections.append((current_section, current_content))
                    current_content = []
                
                current_section = text
            # Add content to current section
            elif current_section is not None:
                # Add paragraph text
                if text:  # Even if it's just a number
                    current_content.append(text)
                
                # Check for form fields or special formatting
                for run in para.runs:
                    if run.bold or run.italic or run.underline:
                        if run.text.strip() and run.text.strip() not in text:
                            current_content.append(f"[FORMATTED: {run.text.strip()}]")
        
        # Add the last section
        if current_section:
            sections.append((current_section, current_content))
        
        # Now integrate tables with sections based on their relative positions
        # For simplicity, we'll associate each table with the section that precedes it
        if tables_info:
            table_index = 0
            for section_index, (section, content) in enumerate(sections):
                # Add regular content
                section_text = "\n".join(content).strip() if content else ""
                
                # Check if there are tables to add to this section
                # For this simple approach, we'll distribute tables evenly among sections
                tables_per_section = len(tables_info) / len(sections)
                start_table = int(section_index * tables_per_section)
                end_table = int((section_index + 1) * tables_per_section)
                
                section_tables = tables_info[start_table:end_table]
                
                # Add tables to section content
                for _, table_text in section_tables:
                    if section_text:
                        section_text += f"\n\nTABLE:\n{table_text}"
                    else:
                        section_text = f"TABLE:\n{table_text}"
                
                # Add to final output
                extracted_data["content"].append({
                    "section": section,
                    "text": section_text
                })
        else:
            # No tables, just add the sections with their content
            for section, content in sections:
                extracted_data["content"].append({
                    "section": section,
                    "text": "\n".join(content).strip() if content else ""
                })
        
        return extracted_data
    
    except Exception as e:
        logger.error(f"Error parsing DOCX file: {str(e)}")
        raise Exception(f"Error parsing DOCX file: {str(e)}")

if __name__ == "__main__":
    docx_path = "sample.docx"
    extracted_data = extract_sections_from_docx(docx_path)
    
    print("\nExtracted Sections and Content:")
    for item in extracted_data["content"]:
        print(f"\nSection: {item['section']}")
        print(f"Content: {item['text'] if item['text'] else '[EMPTY CONTENT]'}")
        print("\n---\n")
